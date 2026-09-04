#!/usr/bin/env python3
"""Build the editable Digital Delta fair report from the reviewed Markdown manuscript."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FINAL_REPORT.md"
OUTPUT = ROOT / "output" / "docx" / "digital-delta-report.docx"

INK = "102A43"
TEAL = "087F8C"
DEEP_TEAL = "075D67"
PALE_TEAL = "E8F4F5"
PALE_BLUE = "EEF4F8"
PALE_GOLD = "FFF5D9"
MID_GRAY = "607D8B"
LIGHT_GRAY = "D9E2EC"
WHITE = "FFFFFF"
BLACK = RGBColor(0x10, 0x2A, 0x43)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LIGHT_GRAY, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("DIGITAL DELTA  •  ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x60, 0x7D, 0x8B)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 31, INK, 0, 9),
        ("Subtitle", 16, TEAL, 0, 14),
        ("Heading 1", 20, INK, 15, 7),
        ("Heading 2", 14, DEEP_TEAL, 12, 5),
        ("Heading 3", 11.5, INK, 9, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"].font.size = Pt(9.2)
    styles["List Number"].font.name = "Aptos"
    styles["List Number"].font.size = Pt(9.2)

    for section in doc.sections:
        section.header_distance = Inches(0.25)
        section.footer_distance = Inches(0.27)
        footer = section.footer.paragraphs[0]
        add_page_number(footer)

    props = doc.core_properties
    props.title = "Digital Delta technical and market report"
    props.subject = "Offline disaster logistics and verified relief delivery in Bangladesh"
    props.author = "Touhidul Alam Seyam"
    props.keywords = "Bangladesh, disaster logistics, offline first, mesh, relief, innovation"
    props.comments = "Prepared for Bangladesh Innovation Fair 2026"


def set_alt_text(inline_shape, description: str) -> None:
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description[:120])


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([props, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|`[^`]+`|https?://\S+)")


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Aptos Mono"
            run.font.size = Pt(8.6)
            run.font.color.rgb = RGBColor.from_string(DEEP_TEAL)
        else:
            clean = token.rstrip(".,")
            add_hyperlink(paragraph, clean, clean)
            if len(clean) != len(token):
                paragraph.add_run(token[len(clean) :])
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def add_full_image(doc: Document, path: Path, width: float, alt: str, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    set_alt_text(shape, alt)
    add_caption(doc, caption)


def add_phone_pair(doc: Document, left: Path, right: Path, caption: str, left_alt: str, right_alt: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, WHITE, "0")
    for cell, path, alt in ((table.cell(0, 0), left, left_alt), (table.cell(0, 1), right, right_alt)):
        cell.width = Inches(3.4)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = paragraph.add_run().add_picture(str(path), width=Inches(2.45))
        set_alt_text(shape, alt)
    add_caption(doc, caption)


def add_status_table(doc: Document) -> None:
    rows = [
        ("Field core", "8 of 8 modules", "Working code paths and automated evidence"),
        ("Languages", "Bangla and English", "Bundled equally; language change preserves state"),
        ("Offline boundary", "Field phone first", "Laptop, Vercel, and D1 are non-authoritative"),
        ("Honest limits", "Visible release gates", "Physical radio, camera, device metrics, rehearsal"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.25)
    table.columns[1].width = Inches(1.55)
    table.columns[2].width = Inches(4.0)
    headers = ("Area", "Current state", "Evidence boundary")
    for index, label in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, DEEP_TEAL)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_shading(cells[index], PALE_BLUE if row_index % 2 == 0 else WHITE)
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[index].paragraphs[0].add_run(value)
    set_table_borders(table)
    doc.add_paragraph()


def add_stack_table(doc: Document) -> None:
    rows = [
        ("Field", "Kotlin, Jetpack Compose, Room, Android Keystore, Nearby Connections"),
        ("Decision engines", "Deterministic Dijkstra, policy triage, ONNX Runtime"),
        ("Protocol", "Protocol Buffers; gRPC on supported IP links; framed Protobuf nearby"),
        ("Observer", "Go, BoltDB, ordered replay, local Server-Sent Events projection"),
        ("Headquarters", "Next.js 16, React, TypeScript, shadcn/ui, MapLibre, PMTiles"),
        ("Optional archive", "Cloudflare Worker and D1 with a strict presentation allow-list"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.15)
    for index, label in enumerate(("Layer", "Implementation")):
        cell = table.cell(0, index)
        set_cell_shading(cell, INK)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_shading(cells[index], PALE_TEAL if row_index % 2 == 0 else WHITE)
            set_cell_margins(cells[index])
            cells[index].paragraphs[0].add_run(value)
    set_table_borders(table)
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.cell(0, 0).width = Inches(6.95)
    set_cell_shading(band.cell(0, 0), INK)
    set_cell_margins(band.cell(0, 0), top=180, start=220, bottom=180, end=220)
    p = band.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("BANGLADESH INNOVATION FAIR 2026")
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    title = doc.add_paragraph(style="Title")
    title.add_run("Digital Delta")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Offline disaster logistics and verified relief delivery in Bangladesh")
    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(13)
    lead_run = lead.add_run("A field-first system that keeps requests, routes, priorities, and custody records moving when internet and power do not.")
    lead_run.font.size = Pt(12)
    lead_run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    add_full_image(
        doc,
        ROOT / "artifacts/screenshots/command-en/02-command-en-live-overview-1920x1080.png",
        6.92,
        "English headquarters overview with the Sylhet mission map, operational status, route, supply and event panels",
        "The live headquarters is an observer. Field phones remain authoritative and continue without it.",
    )

    metadata = doc.add_table(rows=3, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.LEFT
    metadata.autofit = False
    labels = (("Prepared by", "Touhidul Alam Seyam"), ("Role", "Individual project lead and developer"), ("Report date", "4 September 2026"))
    for row_index, (label, value) in enumerate(labels):
        left, right = metadata.rows[row_index].cells
        left.width = Inches(1.25)
        right.width = Inches(4.85)
        set_cell_shading(left, PALE_TEAL)
        set_cell_margins(left, top=70, bottom=70)
        set_cell_margins(right, top=70, bottom=70)
        lr = left.paragraphs[0].add_run(label.upper())
        lr.bold = True
        lr.font.size = Pt(8)
        lr.font.color.rgb = RGBColor.from_string(DEEP_TEAL)
        rr = right.paragraphs[0].add_run(value)
        rr.font.size = Pt(9)
    set_table_borders(metadata, WHITE, "0")
    doc.add_page_break()


def add_contents(doc: Document, headings: list[str]) -> None:
    doc.add_heading("Contents", level=1)
    intro = doc.add_paragraph("This report combines the operational case, technical design, evidence, market position, and pilot plan. It separates working software from physical release gates.")
    intro.paragraph_format.space_after = Pt(10)
    for number, heading in enumerate(headings, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.first_line_indent = Inches(-0.08)
        p.add_run(f"{number}.  {heading}")
    doc.add_page_break()


def extract_body(lines: list[str]) -> tuple[list[str], list[str]]:
    body_start = next(i for i, line in enumerate(lines) if line == "## Executive summary")
    body = lines[body_start:]
    headings = [line[3:] for line in body if line.startswith("## ")]
    return body, headings


def add_markdown_body(doc: Document, body: list[str]) -> None:
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_buffer))
        paragraph_buffer.clear()

    while index < len(body):
        raw = body[index].rstrip()
        if not raw:
            flush_paragraph()
            index += 1
            continue
        if raw.startswith("## "):
            flush_paragraph()
            heading = raw[3:]
            if heading in {
                "The problem in Bangladesh",
                "Module implementation",
                "Bilingual and accessible operation",
                "Market and alternatives",
                "Demonstration plan",
                "Reference register",
            }:
                doc.add_page_break()
            doc.add_heading(heading, level=1)
            if heading == "Executive summary":
                add_status_table(doc)
            elif heading == "System architecture":
                add_stack_table(doc)
            elif heading == "Bilingual and accessible operation":
                add_phone_pair(
                    doc,
                    ROOT / "artifacts/screenshots/field-bn/03-field-bn-p0-request-1280x2856.png",
                    ROOT / "artifacts/screenshots/field-en/03-field-en-p0-request-1280x2856.png",
                    "The same urgent-request workflow in Bangla and English. Both interfaces ship in the APK.",
                    "Bangla urgent relief request form showing P0 medical priority",
                    "English urgent relief request form showing P0 medical priority",
                )
            index += 1
            continue
        if raw.startswith("### "):
            flush_paragraph()
            heading = raw[4:]
            doc.add_heading(heading, level=2)
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", raw)
        bullet = re.match(r"^-\s+(.*)$", raw)
        if numbered:
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.add_run(f"{numbered.group(1)}.  ")
            add_inline(p, numbered.group(2))
            index += 1
            continue
        if bullet:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            index += 1
            continue
        paragraph_buffer.append(raw.strip())
        index += 1
    flush_paragraph()

    doc.add_page_break()
    doc.add_heading("Selected visual evidence", level=1)
    doc.add_heading("Offline multimodal rerouting", level=2)
    add_phone_pair(
        doc,
        ROOT / "artifacts/screenshots/field-bn/14-field-bn-route-rerouted-1280x2856.png",
        ROOT / "artifacts/screenshots/field-en/14-field-en-route-rerouted-1280x2856.png",
        "A simulated E3 failure removes the truck path and the same on-device engine selects the valid boat route.",
        "Bangla route screen showing offline boat reroute after simulated E3 failure",
        "English route screen showing offline boat reroute after simulated E3 failure",
    )
    doc.add_page_break()
    doc.add_heading("Signed custody and replay rejection", level=2)
    add_phone_pair(
        doc,
        ROOT / "artifacts/screenshots/field-en/17-field-en-pod-verified.png",
        ROOT / "artifacts/screenshots/field-en/18-field-en-pod-replay-rejected.png",
        "The first signed handoff is accepted. Reusing its nonce is rejected without changing the custody chain.",
        "English proof-of-delivery screen showing a verified signed handoff",
        "English proof-of-delivery screen showing replay rejection",
    )
    doc.add_page_break()
    doc.add_heading("Delayed boat replanning", level=2)
    add_phone_pair(
        doc,
        ROOT / "artifacts/screenshots/field-bn/24-field-bn-hybrid-replanned-1280x2856.png",
        ROOT / "artifacts/screenshots/field-en/24-field-en-hybrid-replanned-1280x2856.png",
        "An 18-minute simulated boat delay changes the best rendezvous from R3 to R2 before custody transfer.",
        "Bangla hybrid-fleet screen showing delayed-boat replanning to rendezvous R2",
        "English hybrid-fleet screen showing delayed-boat replanning to rendezvous R2",
    )


def main() -> int:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    body, headings = extract_body(lines)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc, headings)
    add_markdown_body(doc, body)
    doc.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
